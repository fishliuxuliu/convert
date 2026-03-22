"""
文件格式转换核心逻辑
支持: docx, doc, pdf, md, txt 互转
"""
import os
import re
import subprocess
import sys
from pathlib import Path

import docx
import pdfplumber


# ---------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------

def get_available_formats():
    """返回支持的格式列表"""
    return ["md", "docx", "pdf", "txt"]


def get_format_display_name(fmt):
    """格式友好名称"""
    names = {
        "md": "Markdown (.md)",
        "docx": "Word 文档 (.docx)",
        "pdf": "PDF 文件 (.pdf)",
        "txt": "纯文本 (.txt)",
    }
    return names.get(fmt, fmt)


def is_pandoc_available():
    """检查 pandoc 是否可用"""
    try:
        subprocess.run(
            ["pandoc", "--version"],
            capture_output=True,
            text=True,
            timeout=10,
        )
        return True
    except Exception:
        return False


# ---------------------------------------------------------------
# 读取文件内容
# ---------------------------------------------------------------

def read_docx(filepath):
    """读取 docx 文件内容"""
    try:
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"读取 docx 文件失败: {filepath}\n{str(e)}")


def read_pdf(filepath):
    """读取 pdf 文件内容"""
    try:
        texts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        return "\n".join(texts)
    except Exception as e:
        raise RuntimeError(f"读取 PDF 文件失败: {filepath}\n{str(e)}")


def read_markdown(filepath):
    """读取 md/txt 文件内容"""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        raise RuntimeError(f"读取文件失败: {filepath}\n{str(e)}")


def read_file(filepath):
    """根据扩展名读取文件内容"""
    ext = Path(filepath).suffix.lower()
    if ext in [".docx"]:
        return read_docx(filepath)
    elif ext in [".pdf"]:
        return read_pdf(filepath)
    elif ext in [".md", ".txt"]:
        return read_markdown(filepath)
    elif ext in [".doc"]:
        # .doc 尝试用 pandoc 转成 txt 再读
        try:
            result = subprocess.run(
                ["pandoc", filepath, "-t", "plain"],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                return result.stdout
        except Exception:
            pass
        raise RuntimeError(f".doc 文件需要 pandoc 支持，请确保已安装 pandoc")
    else:
        raise RuntimeError(f"不支持的文件格式: {ext}")


# ---------------------------------------------------------------
# 写入文件
# ---------------------------------------------------------------

def write_markdown(content, filepath):
    """写入 md 文件"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)


def write_txt(content, filepath):
    """写入 txt 文件"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)


def write_docx(content, filepath):
    """写入 docx 文件"""
    doc = docx.Document()
    # 按换行符分段
    for para in content.split("\n"):
        doc.add_paragraph(para)
    doc.save(filepath)


def write_pdf(content, filepath):
    """写入 pdf 文件（使用 pandoc）"""
    # 先写临时 md，再用 pandoc 转 pdf
    tmp_md = filepath + ".tmp.md"
    with open(tmp_md, "w", encoding="utf-8") as f:
        f.write(content)
    try:
        result = subprocess.run(
            [
                "pandoc", tmp_md,
                "-o", filepath,
                "--pdf-engine=xelatex",
                "-V", "mainfont='Microsoft YaHei'",
                "-V", "geometry:margin=1in",
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            # 如果 xelatex 失败，尝试其他引擎
            result = subprocess.run(
                [
                    "pandoc", tmp_md,
                    "-o", filepath,
                    "-t", "html",
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )
    finally:
        if os.path.exists(tmp_md):
            os.remove(tmp_md)

    if result.returncode != 0 and not os.path.exists(filepath):
        raise RuntimeError(f"PDF 生成失败:\n{result.stderr}")


# ---------------------------------------------------------------
# 格式检测
# ---------------------------------------------------------------

def detect_format(filepath):
    """根据文件扩展名判断格式"""
    ext = Path(filepath).suffix.lower()
    table = {
        ".docx": "docx",
        ".doc": "doc",
        ".pdf": "pdf",
        ".md": "md",
        ".markdown": "md",
        ".txt": "txt",
    }
    return table.get(ext, None)


# ---------------------------------------------------------------
# 核心转换
# ---------------------------------------------------------------

def convert_file(src_path, target_format, callback=None):
    """
    转换单个文件

    参数:
        src_path: 源文件路径
        target_format: 目标格式 (md/docx/pdf/txt)
        callback: 进度回调 (current, total, filename)

    返回:
        (成功数, 失败数, 结果列表)
    """
    src_path = str(src_path)
    src_format = detect_format(src_path)

    if src_format is None:
        return [(False, src_path, f"不支持的格式: {src_path}")]

    if src_format == target_format:
        return [(False, src_path, "源格式与目标格式相同，跳过")]

    # 读取内容
    try:
        content = read_file(src_path)
    except Exception as e:
        return [(False, src_path, str(e))]

    # 构造输出文件名
    base_name = Path(src_path).stem  # 无扩展名的文件名
    src_dir = Path(src_path).parent
    output_name = base_name + "." + target_format
    output_path = src_dir / output_name

    # 重名检测：自动编号
    counter = 1
    while output_path.exists():
        output_name = f"{base_name}_{counter}.{target_format}"
        output_path = src_dir / output_name
        counter += 1

    # 写入
    try:
        if target_format == "md":
            write_markdown(content, str(output_path))
        elif target_format == "txt":
            write_txt(content, str(output_path))
        elif target_format == "docx":
            write_docx(content, str(output_path))
        elif target_format == "pdf":
            write_pdf(content, str(output_path))
        else:
            return [(False, src_path, f"不支持的目标格式: {target_format}")]

        return [(True, str(output_path), "转换成功")]

    except Exception as e:
        # 清理可能产生的垃圾文件
        if output_path.exists():
            try:
                os.remove(str(output_path))
            except Exception:
                pass
        return [(False, src_path, str(e))]


def batch_convert(file_list, target_format, callback=None):
    """
    批量转换文件

    参数:
        file_list: 文件路径列表
        target_format: 目标格式
        callback: 回调函数 (current, total, filename, status)

    返回:
        (成功列表, 失败列表)
    """
    successes = []
    failures = []
    total = len(file_list)

    for i, filepath in enumerate(file_list):
        results = convert_file(filepath, target_format)
        success, path, msg = results[0]
        if callback:
            callback(i + 1, total, os.path.basename(filepath), success, msg)
        if success:
            successes.append(path)
        else:
            failures.append((path, msg))

    return successes, failures
